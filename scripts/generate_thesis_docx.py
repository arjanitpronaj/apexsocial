# -*- coding: utf-8 -*-
"""Generate TemaDiplomes.docx aligned with current ApexSocial stack."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

# WARNING: Do not run this on TemaDiplomes.docx — it OVERWRITES the full thesis.
# Use scripts/patch_thesis_inplace.py for surgical updates only.
OUT = Path(__file__).resolve().parents[1] / "TemaDiplomes_GENERATED_SHORT.docx"


def set_normal_style(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)


def add_title_page(doc):
    for text, size, bold in [
        ("UNIVERSITETI AAB", 14, True),
        ("Fakulteti i Shkencave Kompjuterike", 12, False),
        ("Programi i Studimeve: Siguri Kibernetike", 12, False),
        ("", 12, False),
        ("PUNIM DIPLOME", 14, True),
        ("", 12, False),
        (
            "ApexSocial: Rrjet social me moderim të përmbajtjes "
            "të bazuar në inteligjencë artificiale dhe komunikim në kohë reale",
            13,
            True,
        ),
        ("", 12, False),
        ("", 12, False),
        ("Përgatitur nga: [Emri i studentit]", 12, False),
        ("Mentor: [Emri i mentorit]", 12, False),
        ("Prishtinë, 2026", 12, False),
    ]:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = bold
            run.font.size = Pt(size)
            run.font.name = "Times New Roman"
    doc.add_page_break()


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(1.25)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def build():
    doc = Document()
    set_normal_style(doc)
    add_title_page(doc)

    add_heading(doc, "Përmbledhje", 1)
    add_para(
        doc,
        "Ky punim diplome paraqet projektin ApexSocial — një platformë rrjeti social "
        "e zhvilluar në mjedisin XAMPP (PHP 8, MySQL 8), e cila integron një shërbim "
        "moderimi të përmbajtjes të bazuar në Python (Flask, scikit-learn) dhe një "
        "shtresë komunikimi në kohë reale përmes WebSocket-it nativ (porti 8080) "
        "dhe urës HTTP push (porti 8081). Qëllimi i sistemit është të lejojë "
        "publikimin e postimeve dhe komenteve vetëm pas vlerësimit automatik të "
        "rrezikut, duke klasifikuar përmbajtjen në dy kategori përfundimtare: "
        "ALLOWED (e lejuar) dhe FORBIDDEN (e ndaluar). Moderimi kombinon modele "
        "makinerike (TF-IDF me regresion logjistik), lista fjalësh kyçe, analizë "
        "URL-esh, kontroll integriteti Unicode, rregulla kontekstuale gjuhësore "
        "dhe, opsionalisht, vlerësim semantik me transformer. Sistemi mbështet "
        "mësimin inkremental dhe trajnimin periodik, si dhe feedback nga "
        "administratori për përmirësimin e modeleve. Rezultati është një "
        "arkitekturë hibride që ndan përgjegjësitë: shtresa e prezantimit dhe "
        "ruajtjes së të dhënave në PHP/MySQL, shtresa e inteligjencës në Python, "
        "dhe shtresa e njoftimeve në kohë reale në procesin ws_server.py."
    )

    add_heading(doc, "Abstrakt (anglisht)", 1)
    add_para(
        doc,
        "This thesis presents ApexSocial, a social network platform built on "
        "XAMPP (PHP, MySQL) with a Python-based content moderation microservice "
        "(Flask, scikit-learn) and native WebSocket real-time notifications. "
        "Content is classified into ALLOWED or FORBIDDEN verdicts using a hybrid "
        "pipeline combining machine learning, keyword heuristics, URL analysis, "
        "Unicode integrity checks, and optional semantic scoring. The system "
        "supports incremental learning, scheduled retraining, and administrator "
        "feedback for continuous model improvement."
    )

    add_heading(doc, "Hyrje", 1)
    add_para(
        doc,
        "Rritja e përdorimit të rrjeteve sociale ka sjellë nevojë për mekanizma "
        "automatikë të moderimit të përmbajtjes, veçanërisht për gjuhë ofenduese, "
        "ngacmuese dhe materiale mashtruese (phishing, scam). Moderimi vetëm "
        "manual nuk është i qëndrueshëm për volum të lartë postimesh. ApexSocial "
        "adresohet si zgjidhje e integruar: përdoruesi shkruan postim ose koment, "
        "sistemi vlerëson tekstin para publikimit, dhe vetëm përmbajtja e "
        "klasifikuar si e sigurt ruhet në bazën e të dhënave. Ky punim përshkruan "
        "arkitekturën aktuale të implementuar në projektin rojë (repository), "
        "duke u distancuar nga versionet e hershme që e vendosnin SignalR (C#) "
        "si shtresë kryesore realtime — në implementimin final, komunikimi "
        "në kohë reale realizohet me ws_server.py (Python, websockets + aiohttp)."
    )

    add_heading(doc, "Qëllimet dhe objektivat", 1)
    add_bullets(
        doc,
        [
            "Të projektohet dhe implementohet një rrjet social funksional me autentifikim, postime, komente dhe njoftime.",
            "Të integrohet moderim automatik i përmbajtjes përmes API-së ML në portin 5000.",
            "Të sigurohet parakontroll në klient (browser) me pritje 10 sekondash pas ndërprerjes së shkrimit, para aktivizimit të butonit Post.",
            "Të implementohet komunikim në kohë reale për preview të moderimit dhe njoftime (WebSocket :8080, HTTP push :8081).",
            "Të mbështetet përmirësimi i vazhdueshëm i modeleve përmes logimit të mostrave, trajnimit inkremental dhe feedback-ut të administratorit.",
            "Të respektohen parimet e sigurisë kibernetike: validim hyrjesh, rate limiting, autentifikim i lidhjes WebSocket me token HMAC.",
        ],
    )

    add_heading(doc, "Përshkrimi i problemit", 1)
    add_para(
        doc,
        "Problemi kryesor është klasifikimi i besueshëm i tekstit të lirë (free text) "
        "në kategoritë: e sigurt (ALLOWED) ose e dëmshme (FORBIDDEN), me nënkategori "
        "hate_speech dhe phishing_scam. Sfida teknike përfshijnë: (1) evitimin e "
        "fals pozitivëve për tekst neutral; (2) zbulimin e përpjekjeve për të "
        "anashkaluar filtrin (Unicode i çuditshëm, karaktere të fshehura); "
        "(3) sinkronizimin e vendimit në browser, server PHP dhe shërbimin ML; "
        "(4) ruajtjen e performancës kur numri i përdoruesve rritet. Versioni i "
        "vjetër i dokumentacionit përmendte verdictin REVIEW dhe radhë "
        "moderimi automatike për postime kufitare — në versionin aktual, "
        "borderline_review_only është false dhe vendimi është binar: nëse "
        "combined_score ≥ 52%, përmbajtja konsiderohet FORBIDDEN."
    )

    add_heading(doc, "Arkitektura e sistemit", 1)
    add_heading(doc, "Pamje e përgjithshme", 2)
    add_para(
        doc,
        "Arkitektura ndahet në katër shtresa kryesore. (1) Klienti (HTML, CSS, "
        "JavaScript): assets/js/app.js për formularin e postimit dhe "
        "assets/js/realtime.js për WebSocket. (2) Aplikacioni web (PHP): "
        "includes/config.php, includes/ajax.php, index.php, includes/sanitize.php. "
        "(3) Baza e të dhënave (MySQL, utf8mb4): përdorues, postime, komente, "
        "content_analysis, njoftime. (4) Shërbimet Python në ml_api/: api.py "
        "(Flask/Waitress, port 5000), ws_server.py (WebSocket 8080, push 8081), "
        "train_model.py, scheduler.py, online_learning.py. Konstantet kryesore "
        "në PHP: ML_API_URL = http://127.0.0.1:5000; BACKEND_URL drejton te "
        "push API në portin 8081 përmes includes/realtime.php."
    )

    add_heading(doc, "Diagrami logjik i rrjedhës", 2)
    add_para(
        doc,
        "Rrjedha e moderimit të postimit: Përdoruesi shkruan në #post-content → "
        "app.js çaktivizon Post dhe nis numërimin 10s (reset në çdo input) → "
        "pas 10s pa tastierë, dërgohet kërkesë preview (WebSocket ose AJAX "
        "moderate_content) → Python analyze() kthen JSON me verdict → UI "
        "aktivizon Post vetëm në ALLOWED → në submit, index.php thërret përsëri "
        "moderateContent() në server (vendim përfundimtar) → ruajtje në MySQL "
        "ose mesazh gabimi nëse FORBIDDEN. Kjo dyfishim kontroll (klient + server) "
        "është e nevojshme për siguri, pasi klienti mund të manipulohet."
    )

    add_heading(doc, "Shërbimi i moderimit (ML API)", 1)
    add_heading(doc, "Pipeline i analizës", 2)
    add_para(
        doc,
        "Endpoint-i kryesor është POST /analyze. Hapat në api.py: (1) validate_request "
        "dhe security_reject (rate limit, madhësi trupi, spam i përsëritur); "
        "(2) prepare_text nga text_integrity.py (normalizim Unicode, zbulim "
        "manipulimesh); (3) keyword_match për hate dhe scam; (4) ml_features dhe "
        "parashikim me pipeline.pkl (hate) dhe scam_pipeline.pkl (scam); "
        "(5) apply_context_to_probs për rregulla gjuhësore; (6) url_scam_boost; "
        "(7) llogaritje combined_score dhe krahasim me pragun THRESH_PCT (52%); "
        "(8) kthim ALLOWED ose FORBIDDEN. Nëse integriteti dështon, kthehet "
        "verdict REJECTED (jo e njëjta me FORBIDDEN — refuzim teknik i kërkesës)."
    )

    add_heading(doc, "Modelet dhe trajnimi", 2)
    add_para(
        doc,
        "Modelet përdorin TfidfVectorizer me analyzer='char_wb' (n-grama "
        "karakteresh, rezistent ndaj variacioneve të shkrimit) dhe "
        "LogisticRegression për klasifikim binar. Skedarët ruhen në "
        "ml_api/models/pipeline.pkl dhe scam_pipeline.pkl. Trajnimi offline "
        "bëhet me train_model.py; mostrat grumbullohen në user_inputs.jsonl. "
        "online_learning.py regjistron mostra pas çdo analize FORBIDDEN dhe "
        "mund të nisë ri-trajnim inkremental kur plotësohen kushtet "
        "(retrain_min_samples, retrain_every_n). scheduler.py planifikon "
        "trajnim inkremental natën (03:00 UTC) dhe trajnim të plotë javor "
        "(të dielën 04:00 UTC). Endpoint POST /feedback pranon korrigjime nga "
        "admin/queue.php (approve/reject) me peshë të shtuar për mësim."
    )

    add_heading(doc, "Komunikimi në kohë reale", 1)
    add_para(
        doc,
        "ws_server.py zëvendëson SignalR-in e vjetër (Backend/Program.cs në C# "
        "mbetet në depo si legacy, por nuk është shtresa autoritative). "
        "Klientët lidhen në ws://host:8080 me mesazh join që përfshin user_id "
        "dhe wsToken (HMAC me APEX_WS_KEY / WS_SECRET) — roli admin nuk "
        "deklarohet nga klienti. PHP dërgon njoftime përmes POST /api/push në "
        "portin 8081 me REALTIME_PUSH_KEY. Mesazhet përfshijnë notification, "
        "moderation_result, queue_update, banned. realtime.js menaxhon "
        "rilidhjen, pong timeout dhe limit rate për preview (~1.2s)."
    )

    add_heading(doc, "Shtresa PHP dhe baza e të dhënave", 1)
    add_para(
        doc,
        "Funksioni moderateContent() në config.php bën HTTP POST te /analyze "
        "me tekst, user_id dhe content_type. mlVerdictBlocks() kthen true vetëm "
        "për FORBIDDEN. ajax.php ekspozon action=moderate_content për AJAX. "
        "index.php aplikon moderimin përsëri në submit të formularit. "
        "sanitize.php dhe kontroll XSS në ajax.php mbrojnë nga injektimi. "
        "Tabela content_analysis ruan metadata ML (verdict, category, "
        "harmful_prob). Radha admin (admin/queue.php) shërben për raportime "
        "manuale dhe korrigjim feedback, jo për pritje automatike REVIEW."
    )

    add_heading(doc, "Siguria", 1)
    add_bullets(
        doc,
        [
            "Rate limiting në ML API (security.rate_limit_analyze, flood_max_requests).",
            "Validim gjatësi teksti dhe madhësi body (max_text_length, max_body_bytes).",
            "Integriteti Unicode dhe refuzimi i modeleve të dyshimta (text_integrity.py).",
            "Token HMAC për join WebSocket; çelësi APEX_WS_KEY në mjedis/server.",
            "Sesione PHP, role user/admin, kontroll ban për përdorues të bllokuar.",
            "CORS i kufizuar në origjina lokale në config.json.",
        ],
    )

    add_heading(doc, "Implementimi në mjedisin e zhvillimit", 1)
    add_para(
        doc,
        "Për ekzekutim lokal: (1) XAMPP — Apache dhe MySQL; import database.sql; "
        "(2) python api.py në ml_api/ (port 5000); (3) python ws_server.py "
        "(portet 8080 dhe 8081); (4) opsional python scheduler.py për trajnime "
        "të planifikuara. Frontend-i aksesohet në http://localhost/apexsocial. "
        "Navbar ekspozon APEX_WS_URL për lidhjen WebSocket nga browser-i."
    )

    add_heading(doc, "Vlerësimi dhe diskutimi", 1)
    add_para(
        doc,
        "Sistemi vlerësohet sipas: saktësisë së klasifikimit në mostra test "
        "(hate speech, scam, tekst neutral shqip/anglisht); kohës së përgjigjes "
        "të /analyze; qëndrueshmërisë së lidhjes WebSocket; dhe përvojës së "
        "përdoruesit (pritja 10s redukton thirrje të panevojshme ML gjatë "
        "shkrimit). Kufizimet: modelet varen nga cilësia e të dhënave të "
        "trajnimit; gjuhët e rralla mund të kenë performancë më të ulët; "
        "semantic_scorer kërkon burime shtesë kur është aktiv. Për punime "
        "të ardhshme rekomandohet metrikë F1 në set validimi dhe testim "
        "ngarkese (load testing) për ML API dhe WebSocket."
    )

    add_heading(doc, "Përfundime", 1)
    add_para(
        doc,
        "ApexSocial demonstron integrimin praktik të një aplikacioni web klasik "
        "(PHP/MySQL) me shërbim ML modern (Python/sklearn) dhe realtime "
        "(WebSocket nativ). Moderimi binar ALLOWED/FORBIDDEN thjeshton "
        "përvojën e përdoruesit dhe logjikën e serverit. Arkitektura modulare "
        "lejon zhvillim dhe deploy të pavarur të komponentëve ML. Ky dokument "
        "reflekton gjendjen e projektit rojë në vitin 2026 dhe zëvendëson "
        "përshkrimet e vjetra të bazuara kryesisht në C# SignalR dhe verdictin "
        "REVIEW si rrjedhë kryesore publikimi."
    )

    add_heading(doc, "Literatura dhe burime", 1)
    add_bullets(
        doc,
        [
            "Pedregosa et al., scikit-learn: Machine Learning in Python, JMLR 2011.",
            "RFC 6455 — The WebSocket Protocol.",
            "OWASP — Input Validation and XSS Prevention Cheat Sheets.",
            "Flask Documentation — https://flask.palletsprojects.com/",
            "PHP Manual — PDO, Sessions, Security.",
        ],
    )

    add_heading(doc, "Shtojca A — Tabela e porteve dhe shërbimeve", 1)
    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    hdr = ["Shërbimi", "Porti", "Teknologjia"]
    for i, h in enumerate(hdr):
        table.rows[0].cells[i].text = h
    rows = [
        ("Web + PHP", "80", "Apache (XAMPP)"),
        ("MySQL", "3306", "MySQL 8"),
        ("ML API", "5000", "Flask + Waitress (api.py)"),
        ("WebSocket", "8080", "websockets (ws_server.py)"),
        ("HTTP Push", "8081", "aiohttp (ws_server.py)"),
    ]
    for r, row in enumerate(rows, 1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val

    add_heading(doc, "Shtojca B — Skedarë kryesorë të projektit", 1)
    add_bullets(
        doc,
        [
            "ml_api/api.py, context_scoring.py, text_integrity.py, security.py, online_learning.py",
            "ml_api/ws_server.py, train_model.py, scheduler.py",
            "includes/config.php, ajax.php, realtime.php, sanitize.php",
            "assets/js/app.js, realtime.js",
            "admin/queue.php, index.php",
            "docs/FULL_SYSTEM_AUDIT.md",
        ],
    )

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
