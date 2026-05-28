# -*- coding: utf-8 -*-
"""
Patch TemaDiplomes.docx in place: only outdated C#/SignalR/Singular Architecture refs.
Does NOT shorten or regenerate the document.
"""
from copy import deepcopy
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "TemaDiplomes.docx"
BACKUP = ROOT / "TemaDiplomes_BACKUP_para_patch.docx"


def set_para_text(paragraph, new_text: str) -> None:
    """Replace visible text; keep first run style when possible."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def set_cell_text(cell, new_text: str) -> None:
    if cell.paragraphs:
        set_para_text(cell.paragraphs[0], new_text)
    else:
        cell.text = new_text


# Full paragraph replacements (by index in body)
PARA_FULL = {
    116: (
        "Në kuadër të këtij projekti, synimi kryesor ka qenë analiza, projektimi dhe "
        "implementimi i një sistemi inteligjent për detektimin dhe bllokimin automatik të "
        "përmbajtjeve të dëmshme në rrjetet sociale. Për këtë qëllim është zhvilluar "
        "platforma sociale ApexSocial, e cila integron teknika të Machine Learning në "
        "proceset e publikimit dhe moderimit të përmbajtjes. Sistemi bazohet në një "
        "arkitekturë hibride: PHP (XAMPP) për aplikacionin web dhe MySQL, Python Flask "
        "API (porti 5000) për analizën ML, dhe ws_server.py (WebSocket 8080, HTTP push "
        "8081) për komunikim në kohë reale. Vendimi i moderimit është binar "
        "(ALLOWED / FORBIDDEN). Implementimi përfshin modele trajnimi që kombinojnë "
        "TF-IDF me Logistic Regression, të cilat janë trajnuar mbi dataset-e publike për "
        "identifikimin e gjuhës së urrejtjes dhe përmbajtjeve mashtruese. Rezultatet e "
        "analizës dhe testimit konfirmojnë se sistemi arrin një saktësi mbi 94% në "
        "klasifikimin automatik të përmbajtjes, duke ofruar feedback të menjëhershëm "
        "përpara publikimit."
    ),
    121: (
        "This project presents analyze, design, and implement an intelligent system for "
        "the automated detection and blocking of harmful content on social networks. To "
        "achieve this, the ApexSocial social platform was developed, integrating Machine "
        "Learning techniques into the content publishing and moderation processes. The "
        "system uses a hybrid architecture: PHP (XAMPP) for the web application and "
        "MySQL, a Python Flask ML API (port 5000) for classification, and ws_server.py "
        "(WebSocket 8080, HTTP push 8081) for real-time events. Moderation verdicts are "
        "binary (ALLOWED / FORBIDDEN). The implementation involves training models that "
        "combine TF-IDF with Logistic Regression, trained on public datasets to identify "
        "hate speech and deceptive content. Analysis and testing results indicate that "
        "the system achieves over 94% accuracy in automated content classification, "
        "providing immediate feedback prior to publication."
    ),
    122: (
        "Keywords: content moderation, artificial intelligence, hybrid architecture, "
        "WebSocket, Scikit-learn, TF-IDF, social platform, cybersecurity"
    ),
    193: (
        "Implementimi praktik: Në fund, modeli është integruar në platformën ApexSocial, "
        "ku çdo postim analizohet automatikisht përpara publikimit përmes PHP "
        "(moderateContent) që komunikon me Python Flask API (:5000), me preview në kohë "
        "reale përmes WebSocket (ws_server.py), duke mundësuar moderim para publikimit."
    ),
    224: (
        "Python websockets + aiohttp (ws_server.py) përdoret për komunikim në kohë reale: "
        "WebSocket në portin 8080 për klientët (assets/js/realtime.js) dhe HTTP push në "
        "portin 8081 për njoftime nga PHP (includes/realtime.php)."
    ),
    225: (
        "JavaScript (realtime.js) lidhet me serverin WebSocket nativ; PHP dërgon "
        "njoftime dhe rezultate moderimi përmes POST /api/push me çelës autentifikimi."
    ),
    280: (
        "Aplikacioni është ndërtuar duke përdorur teknologji bazike të zhvillimit web si "
        "HTML për strukturën, CSS për dizajnin dhe JavaScript për ndërveprimin dinamik "
        "(përfshirë app.js dhe realtime.js), ndërsa për backend-in kryesor është përdorur "
        "PHP në XAMPP, i cili menaxhon sesionet, routing-un, AJAX dhe thirrjet te ML API. "
        "Për ruajtjen e të dhënave është përdorur një databazë MySQL (utf8mb4), ku ruhen "
        "informacionet e përdoruesve, postimet, komentet dhe logu i analizave ML."
    ),
    292: (
        "ApexSocial implementon një arkitekturë hibride: PHP orchestron logjikën e "
        "aplikacionit web dhe dërgon tekstin për analizë te Python ML API përmes "
        "moderateContent(); JavaScript ofron parakontroll në browser (pritje 10 sekondash "
        "pas ndërprerjes së shkrimit) para se të aktivizohet publikimi. Vendimi përfundimtar "
        "i moderimit merret nga analyze() në Python dhe verifikohet përsëri në server kur "
        "ruhet postimi."
    ),
    295: "Figura 5: Arkitektura hibride (PHP + Python ML + WebSocket)",
    297: (
        " Python ML API (analyze()) është pika e referencës për verdictin ALLOWED/FORBIDDEN; "
    ),
    300: (
        " WebSocket (ws_server.py) lejon preview moderimi dhe njoftime në kohë reale pa polling."
    ),
    315: (
        "Admin paneli (admin/queue.php) menaxhon përmbajtjet e raportuara, dërgon feedback "
        "te endpoint-i POST /feedback për përmirësimin e modeleve, dhe merr njoftime në kohë "
        "reale përmes WebSocket kur përditësohet radha e moderimit."
    ),
    354: (
        "Kur PHP (moderateContent) ose ws_server.py dërgon tekst te endpoint-i /analyze, "
        "funksioni analyze() zbaton nivele kontrolli të njëpasnjëshme (siguri, integritet "
        "Unicode, fjalëkyçe, ML, kontekst). Secili nivel kontribuon në combined_score; "
        "nëse score ≥ 52%, përmbajtja klasifikohet FORBIDDEN, përndryshe ALLOWED."
    ),
    371: "Integrim të plotë me PHP, WebSocket dhe analizë në kohë reale:",
    373: "PHP dërgon kërkesë POST /analyze (ose preview përmes WebSocket) me:",
    393: (
        "Integrimi i plotë i sistemit ML me arkitekturën hibride (PHP dhe Python), ku secila "
        "shtresë ruan autonominë dhe komunikimi ndodh nëpërmjet kontratave të qarta "
        "ndërfaqeje (REST /analyze, /feedback, WebSocket, HTTP push)."
    ),
}

# Substring replacements applied to any paragraph/table cell (order matters)
SUBSTRING_REPLACEMENTS = [
    ("Arkitekturën Singulatre", "një arkitekturë hibride"),
    ("Singulatre Architecture", "hybrid architecture"),
    ("Singular Architecture", "arkitekturë hibride"),
    ("singular architecture", "hybrid architecture"),
    ("Arkitektura Tre-Shtresor (Singular Architecture)", "Arkitektura hibride (PHP + Python ML + WebSocket)"),
    ("tre-shtresore (PHP, C# dhe Python)", "hibride (PHP dhe Python)"),
    ("C# / ASP.NET Core (Minimal API, .NET 8) përdoret si \"Singular Architecture\", shërbimi qendror që merr kërkesat nga PHP, komunikon me Python ML API dhe menaxhon logjikën e biznesit.", ""),
    ("SignalR përdoret (përmes Microsoft.AspNetCore.SignalR) për komunikim në kohë reale midis komponentëve të sistemit.", ""),
    ("C# ASP.NET Core", "Python ML API"),
    ("C# backend", "PHP / ws_server"),
    ("nga C#", "nga PHP (moderateContent)"),
    ("thirret nga C#", "thirret nga PHP (moderateContent)"),
]

TABLE_REPLACEMENTS = {
    (1, 1, 2): "Analizë e një teksti: thirret nga PHP (moderateContent) para çdo ruajtjeje postimi",
}


def patch_document(doc: Document) -> int:
    changes = 0

    for idx, new_text in PARA_FULL.items():
        if idx < len(doc.paragraphs):
            old = doc.paragraphs[idx].text
            if old != new_text:
                set_para_text(doc.paragraphs[idx], new_text)
                changes += 1

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if not text:
            continue
        new_text = text
        for old, new in SUBSTRING_REPLACEMENTS:
            if old and old in new_text:
                new_text = new_text.replace(old, new)
            elif old == "":
                continue
        # collapse double spaces from removals
        while "  " in new_text:
            new_text = new_text.replace("  ", " ")
        new_text = new_text.strip()
        if new_text != text:
            set_para_text(paragraph, new_text)
            changes += 1

    for (ti, ri, ci), new_text in TABLE_REPLACEMENTS.items():
        if ti < len(doc.tables):
            tbl = doc.tables[ti]
            if ri < len(tbl.rows) and ci < len(tbl.rows[ri].cells):
                cell = tbl.rows[ri].cells[ci]
                if cell.text != new_text:
                    set_cell_text(cell, new_text)
                    changes += 1

    return changes


def main():
    if not DOCX.exists():
        raise SystemExit(f"Missing {DOCX}")

    if not BACKUP.exists():
        import shutil

        shutil.copy2(DOCX, BACKUP)
        print(f"Backup: {BACKUP}")

    doc = Document(DOCX)
    n = patch_document(doc)
    doc.save(DOCX)
    print(f"Patched {n} blocks -> {DOCX}")
    print(f"Paragraphs: {len(doc.paragraphs)}, size bytes: {DOCX.stat().st_size}")


if __name__ == "__main__":
    main()
